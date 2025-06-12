
from docxtpl import DocxTemplate
from docx import Document
from pathlib import Path
import re, tempfile, subprocess
from rules import load_rules, evaluate_condition

TEMPLATE_PATH = Path(__file__).with_name("tr_template.docx")

def filter_docx(doc: Document, rules, context, keep_markers=None, remove_markers=None):
    keep_markers = set(keep_markers or [])
    remove_markers = set(remove_markers or [])
    for cond_name, info in rules['blocos_condicionais'].items():
        marker = info['marcador']
        if marker in remove_markers:
            keep = False
        elif marker in keep_markers:
            keep = True
        else:
            keep = evaluate_condition(info['condicao'], context)
        start_pat = f"[[{info['marcador']}]]"
        end_pat = f"[[END_{info['marcador']}]]"
        inside = False
        to_remove = []
        for p in doc.paragraphs:
            text = p.text
            if start_pat in text:
                inside = True
                p.text = text.replace(start_pat, "")
                if not keep:
                    to_remove.append(p)
                continue
            if end_pat in text:
                inside = False
                p.text = text.replace(end_pat, "")
                if not keep:
                    to_remove.append(p)
                continue
            if inside and not keep:
                to_remove.append(p)
        for p in to_remove:
            p.clear()
    # Remove leftover tags
    for p in doc.paragraphs:
        p.text = re.sub(r"\[\[.*?\]\]", "", p.text)

def render_tr(context, rules=None, keep_markers=None, remove_markers=None):
    if rules is None:
        rules = load_rules()
    tpl = DocxTemplate(TEMPLATE_PATH)
    tpl.render(context)
    tpl.save("temp_render.docx")
    doc = Document("temp_render.docx")
    filter_docx(doc, rules, context, keep_markers, remove_markers)
    doc.save("TR_final.docx")
    # convert to PDF
    try:
        subprocess.run(["docx2pdf", "TR_final.docx", "TR_final.pdf"], check=True)
    except Exception:
        pass
    return "TR_final.docx", "TR_final.pdf"
